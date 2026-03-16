
import re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
 
 
# ──────────────────────────────────────────────────────────────────────────────
# HELPER: parse konten inline ke dalam sebuah paragraph
# Menangani: text node, <strong>/<b>, <em>/<i>, <br>, dan elemen inline lainnya
# ──────────────────────────────────────────────────────────────────────────────
def parse_inline(element, paragraph):
    """
    Iterasi children sebuah elemen dan tambahkan run/break ke `paragraph`.
    Mendukung: text node, <strong>/<b>, <em>/<i>, <br>, elemen inline lain.
    """
    for child in element.children:
        if isinstance(child, NavigableString):
            # FIX BUG 2: text node (NavigableString) sebelumnya di-skip karena
            # kondisi `if child.name:` — sekarang di-handle langsung
            text = str(child)
            if text:
                paragraph.add_run(text)
 
        elif child.name in ('strong', 'b'):
            run = paragraph.add_run(child.get_text())
            run.bold = True
 
        elif child.name in ('em', 'i'):
            run = paragraph.add_run(child.get_text())
            run.italic = True
 
        elif child.name == 'br':
            # FIX BUG 3: <br> sebelumnya tidak di-handle di konteks inline
            paragraph.add_run().add_break()
 
        elif child.name:
            # Elemen inline lain (span, u, dll): ambil teksnya saja
            paragraph.add_run(child.get_text())
 
 
# ──────────────────────────────────────────────────────────────────────────────
# MAIN PARSER: block-level elements
# ──────────────────────────────────────────────────────────────────────────────
def parse_element(element, container):
    if element.name == 'p':
        p = container.add_paragraph()
 
        style = element.get('style', '')
        if 'text-indent' in style:
            p.paragraph_format.first_line_indent = Inches(0.5)
 
        # Gunakan parse_inline agar bold/italic/br di dalam <p> tetap terjaga
        parse_inline(element, p)
 
    elif element.name in ('td', 'th'):
        # FIX BUG 1: <td>/<th> sebelumnya tidak di-handle eksplisit sehingga
        # masuk blok `else` dan hanya iterasi children tanpa membuat paragraph.
        # Akibatnya semua konten inline (<strong>, text, <br>) di dalam <td>
        # hilang karena tidak ada yang membuatkan paragraph/run untuk mereka.
        #
        # Solusi: cek apakah <td> punya child block-level (p, div, table).
        # - Kalau ya → delegasikan ke parse_element rekursif (sudah bisa handle)
        # - Kalau tidak → buat satu paragraph dan isi dengan parse_inline
        has_block = any(
            isinstance(c, Tag) and c.name in ('p', 'div', 'table', 'ul', 'ol')
            for c in element.children
        )
        if has_block:
            for child in element.children:
                if isinstance(child, Tag):
                    parse_element(child, container)
        else:
            p = container.add_paragraph()
            align = element.get('style', '')
            if 'text-align: center' in align or 'text-align:center' in align:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align: right' in align or 'text-align:right' in align:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            parse_inline(element, p)
 
    elif element.name == 'div':
        style = element.get('style', '')
        if 'float: right' in style or 'float:right' in style:
            # Blok tanda tangan: proses semua anak secara inline di satu paragraph
            # FIX BUG 4: sebelumnya hanya cari <p recursive=False> sehingga
            # konten langsung di dalam <div> (teks, <strong>, <br>) hilang.
            p = container.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            parse_inline(element, p)
        else:
            for child in element.children:
                if isinstance(child, Tag):
                    parse_element(child, container)
 
    elif element.name == 'table':
        rows = [tr for tr in element.find_all('tr') if tr.find_parent('table') == element]
        if not rows:
            return
 
        max_cols = max(
            len([td for td in row.find_all(['td', 'th']) if td.find_parent('tr') == row])
            for row in rows
        )
        if max_cols == 0:
            return
 
        table = container.add_table(rows=len(rows), cols=max_cols)
 
        for i, row in enumerate(rows):
            cols = [td for td in row.find_all(['td', 'th']) if td.find_parent('tr') == row]
            for j, col in enumerate(cols):
                cell = table.cell(i, j)
                # Hapus paragraf default kosong yang dibuat Word
                for p in cell.paragraphs:
                    p._element.getparent().remove(p._element)
                # Parse isi <td> ke dalam cell
                parse_element(col, cell)
 
    else:
        # Wrapper generic (body, section, article, dll): lanjutkan rekursi
        for child in element.children:
            if isinstance(child, Tag):
                parse_element(child, container)
 
 
def html_to_docx(html_content, output_filename):
    document = Document()
    soup = BeautifulSoup(html_content, 'html.parser')
    parse_element(soup, document)
    document.save(output_filename)
    print(f"Dokumen berhasil disimpan sebagai {output_filename}")
 
 
def clean_llm_html(content):
    """Membersihkan response LLM agar hanya menyisakan HTML murni."""
    if not content:
        return ""
    content = re.sub(r"```html", "", content, flags=re.IGNORECASE)
    content = re.sub(r"```", "", content)
    content = content.strip()
    content = content.replace("\\n", "\n")
    return content

# Menjalankan fungsi dengan output dari Qwen
html_string = clean_llm_html("```html\n<table border='0' cellpadding='0' cellspacing='0'>\n  <tr>\n    <td><img src='' alt='[image]'/>DEPARTEMEN PEKERJAAN UMUM<br>DIREKTORAT JENDERAL BINA MARGA<br>3. Pattimura No. 20 Gd. Septa Taruna Lt. II Keb-Baru, 12110 Telp (021) 7221950</td>\n  </tr>\n  <tr>\n    <td><h2 style='text-align: center;'>MEMO DINAS</h2><br>No. _________________________</td>\n  </tr>\n  <tr>\n    <td><table border='0' cellpadding='0' cellspacing='0'>\n      <tr><td>Kepada Yth</td><td>: Kepala BPKSDM</td></tr>\n      <tr><td>Dari</td><td>: Direktur Jenderal Bina Marga</td></tr>\n      <tr><td>Perihal</td><td>: Masukan atas pengertian Satu Kesatuan Konstruksi Dari Aspek Pertanggungjawaban Terhadap Kegagalan Bangunan (Jembatan) Sesuai Yang Diamanatkan PP29/2000</td></tr>\n      <tr><td>Tanggal</td><td>: 28 Maret 2007</td></tr>\n      <tr><td>Lampiran</td><td>: 1 (satu) berkas</td></tr>\n      <tr><td>Tembusan</td><td>: 1. Direktur Bina Program<br>2. Direktur Bina Teknik<br>3. Direktur Jalan & Jembatan Wilayah Barat<br>4. Direktur Jalan & Jembatan Wilayah Timur<br>5. Direktur Jalan Bebas Hambatan & Jalan Kota</td></tr>\n    </table></td>\n  </tr>\n  <tr>\n    <td><p>Mengulang Memo Dinas kami No: 09/MD/D/2007 tanggal 12 Januari 2007 perihal tersebut di atas yang ditujukan kepada Bapak, maka kami mohon kiranya Bapak dapat segera memberikan tanggapan atas Pengertian Satu Kesatuan Konstruksi dari aspek pertanggungjawaban terhadap kegagalan bangunan (jembatan) sesuai yang diamanatkan UU18/1999 & PP29/2000, mengingat kebutuhan pembangunan jembatan yang meningkat dan proses pengadaan barang/jasa yang memerlukan penyelesaian segera.</p><p style='text-indent: 40px;'>Demikian kami sampaikan, atas perhatian dan kerjasamanya diucapkan terima kasih.</p></td>\n  </tr>\n  <tr>\n    <td><p style='text-align: right;'>Direktur Jenderal Bina Marga</p></td>\n  </tr>\n  <tr>\n    <td><div style='float: right; text-align: center;'>HENDRIANTO_N<br>NIP 110016212</div></td>\n  </tr>\n</table>\n```")

print(html_string)


html_to_docx(html_string, 'hasil_scan_dokumen.docx')