"""
html_to_docx.py  v3.0
======================
Convert an HTML document produced by the OCR pipeline into a .docx file.

v3.0 — all bugs from v2 fixed:
  BUG 1 — Unit confusion (EMU vs twips):
    CONTENT was computed as EMU (from Twips objects), then passed into Twips()
    again → 635× too large. Fix: CONTENT_TWP is plain-integer twips.
    .width on cells uses Twips(n_twips). OOXML w:pos / w:left use n_twips directly.

  BUG 2 — Whitespace runs from HTML indentation:
    NavigableString '\n        1. Item' produced a run with leading newline + spaces.
    Fix: normalize NavigableString text (collapse \n + surrounding whitespace).

  BUG 3 — Duplicate nested-table rows:
    find_all("tr") is recursive → found inner table <tr>s too.
    Fix: iterate only direct-child rows.

Dependencies:
    pip install python-docx beautifulsoup4 lxml
"""

from __future__ import annotations

import re
import copy
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# ─────────────────────────────────────────────────────
# Page geometry — A4, margins 2.5 cm L/R, 2 cm T/B
# ALL measurements below in TWIPS (integer).
# Use Twips(n) only when assigning to python-docx .width / .height properties.
# Pass raw twips integers directly to OOXML w:pos, w:left, w:hanging, etc.
# ─────────────────────────────────────────────────────
PAGE_W_TWP   = 11906   # A4 width  (twips)
PAGE_H_TWP   = 16838
MARGIN_L_TWP = 1418    # 2.5 cm
MARGIN_R_TWP = 1418
MARGIN_T_TWP = 1134    # 2 cm
MARGIN_B_TWP = 1134
CONTENT_TWP  = PAGE_W_TWP - MARGIN_L_TWP - MARGIN_R_TWP   # 9070 twips

# KV alignment tab stops (twips from left text edge)
KV_COLON_TWP = int(CONTENT_TWP * 0.30)   # ≈ 2721 twips — where ":" lands
KV_VALUE_TWP = int(CONTENT_TWP * 0.34)   # ≈ 3083 twips — where value starts


# ─────────────────────────────────────────────────────
# OOXML helpers
# ─────────────────────────────────────────────────────

def _remove_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    tcPr.append(b)


def _add_borders(cell, color="000000", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for s in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"),   "single")
        e.set(qn("w:sz"),    sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        b.append(e)
    tcPr.append(b)


def _cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for s, v in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:w"),    str(v))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _cell_valign(cell, val="top"):
    tcPr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


def _spacing(para, before: int = 0, after: int = 60):
    pPr = para._p.get_or_add_pPr()
    s = OxmlElement("w:spacing")
    s.set(qn("w:before"), str(before))
    s.set(qn("w:after"),  str(after))
    pPr.append(s)


def _indent(para, left: int = 0, hanging: int = 0, first: int = 0):
    """All values in TWIPS."""
    pPr = para._p.get_or_add_pPr()
    i = OxmlElement("w:ind")
    if left:    i.set(qn("w:left"),      str(left))
    if hanging: i.set(qn("w:hanging"),   str(hanging))
    if first:   i.set(qn("w:firstLine"), str(first))
    pPr.append(i)


def _tab_stops(para, stops: list):
    """
    stops = [(pos_twips, align_str), ...]
    pos_twips: integer twips (OOXML w:pos attribute — twips, NOT EMU)
    align_str: 'left' | 'right' | 'center'
    """
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    for pos, align in stops:
        t = OxmlElement("w:tab")
        t.set(qn("w:val"), align)
        t.set(qn("w:pos"), str(int(pos)))   # raw twips
        tabs.append(t)
    pPr.append(tabs)


def _hrule(doc: Document, sz: int = 12, color: str = "000000"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"),   "single")
    bt.set(qn("w:sz"),    str(sz))
    bt.set(qn("w:space"), "1")
    bt.set(qn("w:color"), color)
    b.append(bt)
    pPr.append(b)
    _spacing(p, 0, 140)
    return p


def _align_from_style(style: str) -> WD_ALIGN_PARAGRAPH:
    s = style.replace(" ", "")
    if "text-align:center"  in s: return WD_ALIGN_PARAGRAPH.CENTER
    if "text-align:right"   in s: return WD_ALIGN_PARAGRAPH.RIGHT
    if "text-align:justify" in s: return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _cell_width_twp(cell) -> int:
    """Return cell width in twips (fallback to CONTENT_TWP/2)."""
    w = cell.width
    if w:
        return int(w) // 635    # EMU → twips (1 twip = 635 EMU)
    return CONTENT_TWP // 2


# ─────────────────────────────────────────────────────
# Inline run builder
# ─────────────────────────────────────────────────────

def _normalize_text(raw: str) -> str:
    """
    Normalize HTML whitespace in a NavigableString:
    - Strip leading newline + indentation (HTML source formatting artefact)
    - Replace mid-text newlines with a space
    - Collapse multiple spaces
    """
    if '\n' not in raw:
        return raw
    # Remove leading newline+spaces (indentation artefact from HTML source)
    t = re.sub(r'^\s*\n\s*', '', raw)
    # Replace any remaining newline + surrounding whitespace with a single space
    t = re.sub(r'[ \t]*\n[ \t]*', ' ', t)
    # Collapse multiple spaces
    t = re.sub(r' {2,}', ' ', t)
    return t


def _runs(para, node, bold=False, italic=False, ul=False, fpt=None):
    """Recursively walk an HTML node and add styled runs to `para`."""
    if isinstance(node, NavigableString):
        t = _normalize_text(str(node))
        if not t.strip():
            # Preserve a single inter-element space if it's just whitespace
            if t == ' ':
                para.add_run(' ')
            return
        r = para.add_run(t)
        r.bold      = bold
        r.italic    = italic
        r.underline = ul
        if fpt:
            r.font.size = Pt(fpt)
        return

    name = (node.name or "").lower()
    b  = bold   or name in ("strong", "b")
    i  = italic or name in ("em",     "i")
    u  = ul     or name in ("u",)

    style = node.get("style", "")
    m  = re.search(r"font-size\s*:\s*([\d.]+)pt", style)
    fp = float(m.group(1)) if m else fpt

    if name == "br":
        para.add_run("\n")
        return
    if name == "img":
        para.add_run("[logo]").italic = True
        return

    for ch in node.children:
        _runs(para, ch, bold=b, italic=i, ul=u, fpt=fp)


# ─────────────────────────────────────────────────────
# HTML pre-processing  (fix Phase-3 output before conversion)
# ─────────────────────────────────────────────────────

def _direct_rows(tbl_tag: Tag) -> list:
    """Return only direct-child <tr> rows (not from nested tables)."""
    rows = []
    for el in tbl_tag.children:
        if not isinstance(el, Tag):
            continue
        if el.name == "tr":
            rows.append(el)
        elif el.name in ("tbody", "thead", "tfoot"):
            for sub in el.children:
                if isinstance(sub, Tag) and sub.name == "tr":
                    rows.append(sub)
    return rows


def _is_kv_table(tag: Tag) -> bool:
    return tag.name == "table" and tag.get("border", "1") in ("0", "")


def _first_td_text(tbl: Tag) -> str:
    rows = _direct_rows(tbl)
    if not rows:
        return ""
    tds = rows[0].find_all("td", recursive=False)
    return tds[0].get_text(strip=True) if tds else ""


def preprocess(html: str) -> str:
    """
    Fix two common Phase-3 HTML mistakes:

    1. Orphaned-label pattern:
         <p>Menimbang</p>
         <table border='0'> first td is ":" </table>
       → inject label as first <td> in the table row.

    2. Wrong column order (: | value) → re-order to (empty_label | : | value).
    """
    soup = BeautifulSoup(html, "lxml")
    page = soup.find("div", class_="page") or soup.body or soup
    main = (page.find("div", style=re.compile(r"font-family")) or page)

    # Pass 1: orphaned label merge
    changed = True
    while changed:
        changed = False
        nodes = [n for n in main.children
                 if not (isinstance(n, NavigableString) and not n.strip())]
        for idx, node in enumerate(nodes):
            if not isinstance(node, Tag) or node.name != "p":
                continue
            label = node.get_text(strip=True)
            if not label:
                continue
            # Find next real sibling
            nxt = None
            for j in range(idx + 1, len(nodes)):
                nxt = nodes[j]
                break
            if nxt is None or not isinstance(nxt, Tag):
                continue
            if not _is_kv_table(nxt) or _first_td_text(nxt) != ":":
                continue

            # Inject label into first <td> of each row that starts with ":"
            for ri, tr in enumerate(_direct_rows(nxt)):
                tds = tr.find_all("td", recursive=False)
                if not tds or tds[0].get_text(strip=True) != ":":
                    continue
                new_td = soup.new_tag("td")
                new_td["style"] = (
                    "padding-right:8px; vertical-align:top; white-space:nowrap;"
                )
                if ri == 0:
                    new_td.string = label
                tr.insert(0, new_td)

            node.decompose()
            changed = True
            break

    # Pass 2: wrong column order (first td is ":")
    for tbl in main.find_all("table"):
        if tbl.get("border", "1") not in ("0", ""):
            continue
        for tr in _direct_rows(tbl):
            tds = tr.find_all("td", recursive=False)
            if not tds or tds[0].get_text(strip=True) != ":":
                continue
            colon_td = tds[0]
            rest = tds[1:]
            for td in list(tds):
                td.extract()
            new_lbl = soup.new_tag("td")
            new_lbl["style"] = (
                "padding-right:8px; vertical-align:top; white-space:nowrap;"
            )
            tr.append(new_lbl)
            tr.append(colon_td)
            for td in rest:
                tr.append(td)

    return str(soup)


# ─────────────────────────────────────────────────────
# BR splitter — split a tag's content at <br> into segments
# ─────────────────────────────────────────────────────

def _br_split(tag: Tag) -> list:
    """
    Split a tag's children at <br> elements.
    Returns list of lists-of-nodes, one per visual line.
    Empty/whitespace-only lines are discarded.
    """
    if tag is None:
        return []

    lines: list = [[]]
    for ch in tag.children:
        if isinstance(ch, Tag) and ch.name == "br":
            lines.append([])
        else:
            lines[-1].append(ch)

    # Filter lines that contain visible text
    result = []
    for line_nodes in lines:
        visible = any(
            (isinstance(n, NavigableString) and _normalize_text(str(n)).strip())
            or (isinstance(n, Tag) and n.get_text(strip=True))
            for n in line_nodes
        )
        if visible:
            result.append(line_nodes)
    return result


# ─────────────────────────────────────────────────────
# Converter
# ─────────────────────────────────────────────────────

class HtmlToDocx:

    def __init__(self):
        self.doc = Document()
        sec = self.doc.sections[0]
        # Use Twips() for python-docx property assignment
        sec.page_width    = Twips(PAGE_W_TWP)
        sec.page_height   = Twips(PAGE_H_TWP)
        sec.left_margin   = Twips(MARGIN_L_TWP)
        sec.right_margin  = Twips(MARGIN_R_TWP)
        sec.top_margin    = Twips(MARGIN_T_TWP)
        sec.bottom_margin = Twips(MARGIN_B_TWP)

        nrm = self.doc.styles["Normal"]
        nrm.font.name = "Times New Roman"
        nrm.font.size = Pt(12)
        nrm.paragraph_format.space_before = Pt(0)
        nrm.paragraph_format.space_after  = Pt(3)

        for lvl, pt in [(1, 14), (2, 12), (3, 12), (4, 11)]:
            h = self.doc.styles[f"Heading {lvl}"]
            h.font.name      = "Times New Roman"
            h.font.size      = Pt(pt)
            h.font.bold      = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after  = Pt(3)

    # ── Entry ────────────────────────────────────────────────────────────────

    def convert(self, html: str) -> Document:
        html = preprocess(html)
        soup = BeautifulSoup(html, "lxml")
        page = soup.find("div", class_="page") or soup.body or soup
        main = page.find("div", style=re.compile(r"font-family")) or page
        self._block(main)
        return self.doc

    # ── Block dispatcher ─────────────────────────────────────────────────────

    def _block(self, container):
        for ch in container.children:
            if isinstance(ch, NavigableString):
                t = _normalize_text(str(ch)).strip()
                if t:
                    p = self.doc.add_paragraph(t)
                    _spacing(p)
            elif isinstance(ch, Tag):
                self._tag(ch)

    def _tag(self, tag: Tag):
        name  = tag.name or ""
        style = tag.get("style", "").replace(" ", "")

        if name == "div" and ("border-bottom" in style
                              or ("margin-bottom" in style and tag.find("table"))):
            self._letterhead(tag); return

        if name == "div" and "display:flex" in style:
            self._dual_col(tag); return

        if name == "div" and "float:right" in style:
            self._float_right(tag); return

        if name == "div" and "clear:both" in style:
            return

        if name == "div":
            self._block(tag); return

        if name in ("h1", "h2", "h3", "h4"):
            h = self.doc.add_heading("", level=int(name[1]))
            h.alignment = _align_from_style(style)
            _spacing(h, 120, 60)
            _runs(h, tag)
            return

        if name == "table":
            b = tag.get("border", "0")
            if b and b != "0":
                self._data_table(tag)
            else:
                self._kv_table(tag)
            return

        if name == "ol": self._ol(tag); return
        if name == "ul": self._ul(tag); return

        if name in ("p", "small", "span"):
            self._para(tag); return

        if name == "br":
            p = self.doc.add_paragraph()
            _spacing(p, 0, 0)
            return

        self._block(tag)

    # ── Letterhead ───────────────────────────────────────────────────────────

    def _letterhead(self, div: Tag):
        inner = div.find("table")
        if inner:
            tds = inner.find("tr").find_all("td") if inner.find("tr") else []
            if len(tds) >= 2:
                lw = int(CONTENT_TWP * 0.13)
                rw = CONTENT_TWP - lw

                tbl = self.doc.add_table(rows=1, cols=2)
                tbl.style = "Table Grid"
                lc = tbl.rows[0].cells[0]
                rc = tbl.rows[0].cells[1]
                lc.width = Twips(lw)   # python-docx property → use Twips()
                rc.width = Twips(rw)
                _remove_borders(lc); _remove_borders(rc)
                _cell_valign(lc, "center"); _cell_valign(rc, "center")

                lp = lc.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _spacing(lp, 0, 0)
                lp.add_run("[logo]").italic = True

                first = True
                for el in tds[1].children:
                    if isinstance(el, NavigableString):
                        t = _normalize_text(str(el)).strip()
                        if not t:
                            continue
                        p = rc.paragraphs[0] if first else rc.add_paragraph()
                        first = False
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _spacing(p, 0, 0)
                        p.add_run(t)
                    elif isinstance(el, Tag):
                        if el.name == "br":
                            continue
                        p = rc.paragraphs[0] if first else rc.add_paragraph()
                        first = False
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _spacing(p, 0, 0)
                        _runs(p, el)
                # Remove leading empty paragraph
                if rc.paragraphs and not rc.paragraphs[0].text.strip():
                    e = rc.paragraphs[0]._element
                    e.getparent().remove(e)
        else:
            self._block(div)

        _hrule(self.doc)

    # ── KV table (tab-stop alignment) ────────────────────────────────────────

    def _kv_table(self, tbl_tag: Tag,
                  colon_twp: int = KV_COLON_TWP,
                  value_twp: int = KV_VALUE_TWP):
        """
        Render a borderless KV table using Word tab stops.
        Each row:   LABEL <tab> : <tab> VALUE (first line)
                    [indent]              VALUE (continuation lines)
        """
        for tr in _direct_rows(tbl_tag):
            tds = tr.find_all("td", recursive=False)
            if not tds:
                continue

            # Determine columns
            if len(tds) >= 3:
                lbl_el, val_el = tds[0], tds[2]
            elif len(tds) == 2:
                if tds[0].get_text(strip=True) == ":":
                    lbl_el, val_el = None, tds[1]
                else:
                    lbl_el, val_el = tds[0], tds[1]
            else:
                p = self.doc.add_paragraph()
                _spacing(p, 0, 40)
                _runs(p, tds[0])
                continue

            label_text = lbl_el.get_text(strip=True) if lbl_el else ""

            # Nested KV table inside value cell?
            nested = val_el.find("table") if val_el else None
            if nested:
                # Collect text nodes before the nested table
                intro = []
                for ch in val_el.children:
                    if isinstance(ch, Tag) and ch.name == "table":
                        break
                    intro.append(ch)
                p = self._kv_line(label_text, colon_twp, value_twp)
                for n in intro:
                    _runs(p, n)
                # Render nested table indented by value_twp
                nested_colon = colon_twp + int(CONTENT_TWP * 0.02)
                nested_value = value_twp + int(CONTENT_TWP * 0.02)
                self._kv_table(nested, nested_colon, nested_value)
                continue

            # Split value at <br> for multi-line
            val_lines = _br_split(val_el)

            p = self._kv_line(label_text, colon_twp, value_twp)
            if val_lines:
                for node in val_lines[0]:
                    _runs(p, node)
                for line_nodes in val_lines[1:]:
                    cont = self.doc.add_paragraph()
                    _spacing(cont, 0, 0)
                    # Continuation indent: align with value column
                    _indent(cont, left=value_twp)
                    for node in line_nodes:
                        _runs(cont, node)

        sp = self.doc.add_paragraph()
        _spacing(sp, 0, 80)

    def _kv_line(self, label: str,
                 colon_twp: int = KV_COLON_TWP,
                 value_twp: int = KV_VALUE_TWP):
        """
        Build one KV tab-stop line.
        Tab positions in TWIPS → used as OOXML w:pos attributes directly.
        """
        p = self.doc.add_paragraph()
        _spacing(p, 0, 0)
        # Hanging indent so continuation lines align with value start
        _indent(p, left=value_twp, hanging=value_twp)
        # Tab stops: twips integers → passed to w:pos directly
        _tab_stops(p, [(colon_twp, "left"), (value_twp, "left")])
        if label:
            p.add_run(label)
        p.add_run("\t")
        p.add_run(":")
        p.add_run("\t")
        return p

    # ── Bordered data table ──────────────────────────────────────────────────

    def _data_table(self, tbl_tag: Tag):
        all_trs = _direct_rows(tbl_tag)
        if not all_trs:
            return
        ncols = max(len(r.find_all(["td", "th"], recursive=False)) for r in all_trs)
        if ncols == 0:
            return
        col_w_twp = CONTENT_TWP // ncols

        tbl = self.doc.add_table(rows=0, cols=ncols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        thead_ids = set()
        thead = tbl_tag.find("thead")
        if thead:
            for tr in _direct_rows(thead):
                thead_ids.add(id(tr))

        for ri, tr in enumerate(all_trs):
            ch_cells = tr.find_all(["td", "th"], recursive=False)
            row   = tbl.add_row()
            is_hdr = id(tr) in thead_ids or ri == 0
            for ci, ch in enumerate(ch_cells[:ncols]):
                cell = row.cells[ci]
                cell.width = Twips(col_w_twp)
                _add_borders(cell)
                _cell_margins(cell)
                p = cell.paragraphs[0]
                _spacing(p, 0, 0)
                if "text-align:center" in ch.get("style", "").replace(" ", ""):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _runs(p, ch, bold=is_hdr)

        sp = self.doc.add_paragraph()
        _spacing(sp, 0, 120)

    # ── Lists ────────────────────────────────────────────────────────────────

    def _ol(self, ol: Tag):
        for li in ol.find_all("li", recursive=False):
            p = self.doc.add_paragraph(style="List Number")
            _spacing(p, 0, 30)
            _runs(p, li)

    def _ul(self, ul: Tag):
        for li in ul.find_all("li", recursive=False):
            p = self.doc.add_paragraph(style="List Bullet")
            _spacing(p, 0, 30)
            _runs(p, li)

    # ── Paragraph ────────────────────────────────────────────────────────────

    def _para(self, tag: Tag):
        style = tag.get("style", "")
        p = self.doc.add_paragraph()
        p.alignment = _align_from_style(style.replace(" ", ""))
        _spacing(p)
        if "text-indent:40px" in style.replace(" ", ""):
            _indent(p, first=720)
        _runs(p, tag)

    # ── Float-right ──────────────────────────────────────────────────────────

    def _float_right(self, div: Tag):
        lw_twp = int(CONTENT_TWP * 0.42)
        rw_twp = CONTENT_TWP - lw_twp

        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        lc = tbl.rows[0].cells[0]
        rc = tbl.rows[0].cells[1]
        lc.width = Twips(lw_twp)
        rc.width = Twips(rw_twp)
        _remove_borders(lc); _remove_borders(rc)
        _cell_valign(lc, "top"); _cell_valign(rc, "top")
        _spacing(lc.paragraphs[0], 0, 0)
        self._fill_cell(rc, div, WD_ALIGN_PARAGRAPH.CENTER, rw_twp)
        sp = self.doc.add_paragraph()
        _spacing(sp, 0, 60)

    # ── Dual-column flex ─────────────────────────────────────────────────────

    def _dual_col(self, div: Tag):
        child_divs = [c for c in div.children
                      if isinstance(c, Tag) and c.name == "div"]
        if len(child_divs) < 2:
            self._block(div)
            return

        lw_twp = int(CONTENT_TWP * 0.44)
        rw_twp = CONTENT_TWP - lw_twp

        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        lc = tbl.rows[0].cells[0]
        rc = tbl.rows[0].cells[1]
        lc.width = Twips(lw_twp)
        rc.width = Twips(rw_twp)
        _remove_borders(lc); _remove_borders(rc)
        _cell_valign(lc, "top"); _cell_valign(rc, "top")

        self._fill_cell(lc, child_divs[0], WD_ALIGN_PARAGRAPH.LEFT,   lw_twp)
        self._fill_cell(rc, child_divs[1], WD_ALIGN_PARAGRAPH.CENTER, rw_twp)

        sp = self.doc.add_paragraph()
        _spacing(sp, 0, 80)

    def _fill_cell(self, cell, div: Tag,
                   align: WD_ALIGN_PARAGRAPH, cell_width_twp: int):
        """Render a div's children into a table cell."""
        first = True

        def _p(a=None):
            nonlocal first
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.alignment = a or align
            _spacing(p, 0, 30)
            return p

        for ch in div.children:
            if isinstance(ch, NavigableString):
                t = _normalize_text(str(ch)).strip()
                if t:
                    _runs(_p(), ch)

            elif isinstance(ch, Tag):
                n = ch.name
                if n == "br":
                    p = cell.paragraphs[0] if first else cell.add_paragraph()
                    if first: first = False
                    _spacing(p, 0, 0)
                    continue

                if n == "table":
                    # Nested KV / place-date table → render as KV lines in cell
                    colon_p = int(cell_width_twp * 0.50)
                    value_p = int(cell_width_twp * 0.60)
                    for tr in _direct_rows(ch):
                        tds = tr.find_all("td", recursive=False)
                        if len(tds) >= 3:
                            lbl = tds[0].get_text(strip=True)
                            val = tds[2].get_text(strip=True)
                        elif len(tds) == 2:
                            first_txt = tds[0].get_text(strip=True)
                            lbl = "" if first_txt == ":" else first_txt
                            val = tds[-1].get_text(strip=True)
                        else:
                            lbl = tds[0].get_text(strip=True) if tds else ""
                            val = ""
                        p = cell.paragraphs[0] if first else cell.add_paragraph()
                        if first: first = False
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        _spacing(p, 0, 20)
                        _indent(p, left=value_p, hanging=value_p)
                        _tab_stops(p, [(colon_p, "left"), (value_p, "left")])
                        if lbl: p.add_run(lbl)
                        p.add_run("\t:\t")
                        p.add_run(val)
                    continue

                if n == "p":
                    _runs(_p(), ch)
                    continue

                _runs(_p(), ch)


# ─────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────

def html_to_docx(html: str, output_path: str) -> str:
    """
    Convert an HTML string (from the OCR pipeline) to a .docx file.

    Args:
        html:        Full HTML page string.
        output_path: Destination .docx path.

    Returns:
        Resolved output path string.
    """
    conv = HtmlToDocx()
    doc  = conv.convert(html)
    out  = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python html_to_docx.py input.html output.docx")
        sys.exit(1)
    src  = Path(sys.argv[1]).read_text(encoding="utf-8")
    dest = html_to_docx(src, sys.argv[2])
    print(f"Saved → {dest}")